"""Synchronizuje pola dynamiczne wzorów PSF z aktualnymi formularzami projektu."""

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "public" / "wzory"


def ustaw_akapit(paragraph, tekst: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = tekst
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(tekst)


def ustaw_akapit_po_poczatku(document, poczatek: str, tekst: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(poczatek):
            ustaw_akapit(paragraph, tekst)
            return
    raise RuntimeError(f"Nie znaleziono akapitu: {poczatek}")


def ustaw_komorke(cell, tekst: str) -> None:
    ustaw_akapit(cell.paragraphs[0], tekst)
    for paragraph in cell.paragraphs[1:]:
        ustaw_akapit(paragraph, "")


def aktualizuj_pak1() -> None:
    path = ROOT / "PSF_PAK1_Zgloszenie_potrzeb_szablon.docx"
    doc = Document(path)
    zmiany = {
        "Najwyższy ukończony poziom": "Poziom wykształcenia (ISCED):",
        "{{cb_isced02}}": "{{cb_isced0}} Brak (ISCED 0)    {{cb_isced1}} Podstawowe (ISCED 1)",
        "{{cb_isced34}}": "{{cb_isced2}} Gimnazjalne (ISCED 2)    {{cb_isced3}} Pogimnazjalne / zawodowe (ISCED 3)",
        "{{cb_isced58}} Wyższe": "{{cb_isced4}} Policealne (ISCED 4)    {{cb_isced58}} Wyższe (ISCED 5-8)",
        "☐ Osoba z niepełnosprawnościami": "{{cb_psf_niepelnosprawnosc}} Osoba z niepełnosprawnościami",
        "☒ Nie dotyczy żadnego": "{{cb_psf_brak_statusow_spolecznych}} Nie dotyczy żadnego z powyższych",
        "SUMA PUNKTÓW PREMIUJĄCYCH:": "SUMA PUNKTÓW PREMIUJĄCYCH: {{pkt_psf_premiujace}} / SUMA OGÓŁEM (merytoryczne + premiujące): {{pkt_psf_ogolem}}",
    }
    for poczatek, tekst in zmiany.items():
        ustaw_akapit_po_poczatku(doc, poczatek, tekst)
    doc.save(path)


def aktualizuj_pak2() -> None:
    path = ROOT / "PSF_PAK2_Karta_doradztwa_bilans_szablon.docx"
    doc = Document(path)
    tabela = doc.tables[1]
    nazwy = ["cyfrowe", "jezykowe", "spoleczne", "zawodowe", "zielone"]
    for row_index, nazwa in enumerate(nazwy, start=1):
        for poziom in range(1, 6):
            ustaw_komorke(
                tabela.cell(row_index, poziom),
                f"{{{{cb_psf_{nazwa}_{poziom}}}}}",
            )
    doc.save(path)


def aktualizuj_ipr() -> None:
    path = ROOT / "PSF_B_Indywidualny_Plan_Rozwoju_szablon.docx"
    doc = Document(path)
    tabela = doc.tables[1]
    nazwy = ["cyfrowe", "jezykowe", "spoleczne", "zawodowe", "zielone"]
    for row_index, nazwa in enumerate(nazwy, start=1):
        ustaw_komorke(tabela.cell(row_index, 1), f"{{{{psf_poziom_{nazwa}}}}}")
    doc.save(path)


if __name__ == "__main__":
    aktualizuj_pak1()
    aktualizuj_pak2()
    aktualizuj_ipr()
